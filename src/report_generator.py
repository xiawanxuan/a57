import os
from datetime import datetime
from typing import Dict, List, Optional

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from .config_manager import ConfigManager


class ReportGenerator:
    def __init__(self, config_manager: ConfigManager):
        self.config = config_manager
        self.settings = config_manager.get_report_settings()
        self.global_settings = config_manager.get_global_settings()
        self.doc = None

    def _set_default_font(self, paragraph, size: Optional[int] = None):
        if size is None:
            size = self.settings.font_size
        for run in paragraph.runs:
            run.font.name = self.settings.font_family
            run.font.size = Pt(size)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), self.settings.font_family)

    def _add_heading_cn(self, text: str, level: int = 1):
        h = self.doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.name = self.settings.font_family
            run.font.color.rgb = RGBColor(44, 62, 80)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), self.settings.font_family)
        return h

    def _add_paragraph_cn(self, text: str, bold: bool = False, size: Optional[int] = None):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = self.settings.font_family
        run.font.size = Pt(size if size else self.settings.font_size)
        run.font.bold = bold
        run._element.rPr.rFonts.set(qn("w:eastAsia"), self.settings.font_family)
        return p

    def _add_dataframe_table(self, df: pd.DataFrame, title: Optional[str] = None, max_rows: int = 50):
        if title:
            self._add_paragraph_cn(title, bold=True, size=self.settings.font_size + 1)

        df_display = df.head(max_rows).copy()
        for col in df_display.select_dtypes(include=["float64", "float32"]).columns:
            df_display[col] = df_display[col].apply(lambda x: f"{x:.4f}" if pd.notna(x) else "-")
        df_display = df_display.fillna("-")

        table = self.doc.add_table(rows=len(df_display) + 1, cols=len(df_display.columns))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Light Grid Accent 1"

        for j, col in enumerate(df_display.columns):
            cell = table.rows[0].cells[j]
            cell.text = str(col)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.name = self.settings.font_family
                    run.font.size = Pt(self.settings.font_size - 1)
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), self.settings.font_family)

        for i in range(len(df_display)):
            for j, col in enumerate(df_display.columns):
                cell = table.rows[i + 1].cells[j]
                cell.text = str(df_display.iloc[i, j])
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = self.settings.font_family
                        run.font.size = Pt(self.settings.font_size - 1)
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), self.settings.font_family)

        if len(df) > max_rows:
            self._add_paragraph_cn(f"（表格仅显示前 {max_rows} 行，完整数据见附件 CSV 文件）", size=self.settings.font_size - 2)
        self.doc.add_paragraph()

    def _add_image_safe(self, image_path: str, width: float = 6.0):
        if image_path and os.path.exists(image_path):
            try:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(image_path, width=Inches(width))
            except Exception:
                self._add_paragraph_cn(f"[图表文件: {os.path.basename(image_path)} 未成功嵌入]")

    def generate_report(
        self,
        cleaned_data: Dict[str, pd.DataFrame],
        cleaning_reports: Dict[str, Dict],
        merged_df: pd.DataFrame,
        analysis_results: Dict[str, pd.DataFrame],
        figure_paths: Dict[str, Dict[str, str]],
        comparison_results: Optional[Dict[str, pd.DataFrame]] = None,
        output_filename: str = "paleoclimate_analysis_report.docx",
    ) -> str:
        self.doc = Document()

        self._add_cover_page()
        self._add_toc_placeholder()
        self._add_introduction()
        self._add_data_sources_section(cleaned_data, cleaning_reports)
        self._add_cleaning_methodology()
        self._add_interpolation_methodology()
        self._add_results_section(merged_df, analysis_results, figure_paths)
        self._add_correlation_section(analysis_results)

        if comparison_results:
            self._add_multi_core_comparison_section(comparison_results, figure_paths)

        self._add_conclusion()
        self._add_appendix()

        out_dir = self.config.get_output_dir()
        filepath = os.path.join(out_dir, output_filename)
        self.doc.save(filepath)
        return filepath

    def _add_cover_page(self):
        for _ in range(6):
            self.doc.add_paragraph()

        title_p = self.doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run(self.settings.title)
        run.font.name = self.settings.font_family
        run.font.size = Pt(26)
        run.font.bold = True
        run.font.color.rgb = RGBColor(44, 62, 80)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), self.settings.font_family)

        self.doc.add_paragraph()

        sub_p = self.doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub_p.add_run("——基于冰芯与石笋代用指标的古气候重建")
        run.font.name = self.settings.font_family
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(127, 140, 141)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), self.settings.font_family)

        for _ in range(8):
            self.doc.add_paragraph()

        info_lines = [
            f"编制单位：{self.settings.author}",
            f"报告日期：{datetime.now().strftime('%Y年%m月%d日')}",
            f"数据单位：{self.global_settings.age_unit} / {self.global_settings.temperature_unit} / {self.global_settings.precipitation_unit}",
        ]
        for line in info_lines:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.name = self.settings.font_family
            run.font.size = Pt(12)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), self.settings.font_family)

        self.doc.add_page_break()

    def _add_toc_placeholder(self):
        self._add_heading_cn("目  录", level=1)
        self._add_paragraph_cn("（目录可在 Word 中通过右键→更新域自动生成）")
        self.doc.add_page_break()

    def _add_introduction(self):
        self._add_heading_cn("1  引言", level=1)
        self._add_paragraph_cn(
            "第四纪古气候研究通过分析冰芯、石笋等地质载体中的气候代用指标，"
            "重建万年至百万年尺度的温度与降水演化序列。本报告基于冰芯氧同位素 "
            "（δ¹⁸O）、石笋微量元素（Mg/Ca、Sr/Ca）、粉尘浓度等多源代用指标数据，"
            "经过数据清洗、年代插值对齐、相关性分析与分层可视化，系统展示了研究区域 "
            "内晚第四纪以来的气候波动特征。"
        )
        self._add_paragraph_cn(
            "报告内容涵盖：数据来源与质量评估、处理方法说明、万年尺度时序重建结果、"
            "多指标相关性统计、各地层单位气候特征对比，以及主要结论与展望。"
        )
        self.doc.add_paragraph()

    def _add_data_sources_section(
        self,
        cleaned_data: Dict[str, pd.DataFrame],
        cleaning_reports: Dict[str, Dict],
    ):
        self._add_heading_cn("2  数据来源与清洗质量", level=1)
        self._add_heading_cn("2.1  数据源概况", level=2)

        source_rows = []
        for source_id, df in cleaned_data.items():
            report = cleaning_reports.get(source_id, {})
            source_rows.append({
                "数据源ID": source_id,
                "样本数": len(df),
                "年龄范围(yr BP)": f"{df['age'].min():.0f} ~ {df['age'].max():.0f}" if "age" in df.columns else "-",
                "保留率": f"{report.get('retention_rate', 0):.2%}",
                "剔除异常值": report.get("outliers_total", 0),
            })
        self._add_dataframe_table(pd.DataFrame(source_rows), title="表 1  数据源概况统计")

        self._add_heading_cn("2.2  异常值处理", level=2)
        self._add_paragraph_cn(
            f"本研究采用 Z-Score 阈值法（阈值 = {self.global_settings.outlier_zscore_threshold}σ）"
            "按地层单位分组检测冰芯 δ¹⁸O、石笋 δ¹⁸O、Mg/Ca、粉尘浓度等代用指标中的异常离群值。"
            "对于同时超出该指标在地层单元内 3 倍标准差的数据点，标记为 OUTLIER 并予以剔除。"
        )

        detail_rows = []
        for source_id, report in cleaning_reports.items():
            if "error" in report:
                continue
            for k, v in report.items():
                if k.startswith("outliers_") and k != "outliers_total":
                    detail_rows.append({
                        "数据源": source_id,
                        "指标": k.replace("outliers_", ""),
                        "异常点数": v,
                    })
        if detail_rows:
            self._add_dataframe_table(pd.DataFrame(detail_rows), title="表 2  各指标异常值剔除统计")

    def _add_cleaning_methodology(self):
        self._add_heading_cn("3  数据清洗方法", level=1)
        self._add_paragraph_cn(
            "数据清洗流程包括以下关键步骤："
        )
        steps = [
            "（1）缺失值处理：对 depth、age 关键字段缺失的样本直接剔除；其余数值字段用同列中位数填充。",
            "（2）类型转换：确保所有代用指标列为数值类型，不可解析值设为 NaN。",
            "（3）深度-年龄排序：按地层深度升序、年龄升序排列，保证沉积序列正确。",
            "（4）地层单元标注：根据年龄值自动归属于全新世、晚更新世、中更新世、早更新世四个时代。",
            "（5）分地层异常值剔除：基于 Z-Score/IQR 方法在各地层单元内分别检测并剔除离群点。",
            "（6）重复样本合并：对深度与年龄完全一致的重复测量取均值。",
            "（7）物理量反演：根据数据源配置中的转换系数，由 δ¹⁸O 反演温度（冰芯）与降水（石笋）。",
        ]
        for s in steps:
            self._add_paragraph_cn(s)

    def _add_interpolation_methodology(self):
        self._add_heading_cn("4  年代插值与对齐", level=1)
        self._add_paragraph_cn(
            f"不同代用指标在原始采样中的时间分辨率存在显著差异。为支持多指标联合分析，"
            f"本研究构建统一的年代栅格（分辨率 {self.global_settings.age_grid_resolution} yr），"
            f"采用 {self.global_settings.interpolation_method} 插值方法将各指标序列对齐至公共时间轴。"
        )
        self._add_paragraph_cn(
            "插值完成后，再按照 100 年时间窗进行分箱聚合，以降低高频噪声并突出万年-千年尺度的气候信号。"
            "温度、降水等连续变量取均值，粉尘等受极端值影响的变量取中位数。"
        )

    def _add_results_section(
        self,
        merged_df: pd.DataFrame,
        analysis_results: Dict[str, pd.DataFrame],
        figure_paths: Dict[str, Dict[str, str]],
    ):
        self._add_heading_cn("5  时序重建结果", level=1)

        self._add_heading_cn("5.1  多指标万年时序", level=2)
        self._add_paragraph_cn(
            "图 1 展示了所有启用代用指标在统一年代轴上的多轴时序曲线。"
            "不同颜色与线型分别代表温度、降水、δ¹⁸O、粉尘等指标，背景色块标注了各大地层时代。"
        )
        png = figure_paths.get("multi_axis_timeseries", {}).get("png")
        self._add_image_safe(png, width=6.2)
        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run("图 1  多代用指标万年尺度联合时序曲线")
        r.font.name = self.settings.font_family
        r.font.size = Pt(10)
        r._element.rPr.rFonts.set(qn("w:eastAsia"), self.settings.font_family)

        self._add_heading_cn("5.2  温度与降水波动", level=2)
        self._add_paragraph_cn(
            "图 2 聚焦温度与降水两个核心气候变量的反演结果，直观呈现末次冰期-间冰期旋回中的冷-暖、干-湿变化。"
        )
        png2 = figure_paths.get("temperature_precipitation", {}).get("png")
        self._add_image_safe(png2, width=6.2)
        cap2 = self.doc.add_paragraph()
        cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = cap2.add_run("图 2  万年尺度温度（左轴）与降水（右轴）波动时序")
        r2.font.name = self.settings.font_family
        r2.font.size = Pt(10)
        r2._element.rPr.rFonts.set(qn("w:eastAsia"), self.settings.font_family)

        self._add_heading_cn("5.3  描述性统计", level=2)
        if self.settings.include_summary_statistics and "summary_statistics" in analysis_results:
            self._add_dataframe_table(
                analysis_results["summary_statistics"],
                title="表 3  各代用指标总体描述性统计",
            )

        self._add_heading_cn("5.4  各地层单位分布对比", level=2)
        self._add_paragraph_cn(
            "图 3～图 6 以箱线图形式对比了全新世、晚更新世、中更新世等不同地层单位中主要代用指标的分布差异。"
        )
        img_count = 0
        for key, paths in figure_paths.items():
            if key.startswith("boxplot_") and "png" in paths:
                img_count += 1
                self._add_image_safe(paths["png"], width=5.5)
                cap_p = self.doc.add_paragraph()
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                rr = cap_p.add_run(f"图 {2 + img_count}  {key.replace('boxplot_', '')} 各地层单位分布箱线图")
                rr.font.name = self.settings.font_family
                rr.font.size = Pt(10)
                rr._element.rPr.rFonts.set(qn("w:eastAsia"), self.settings.font_family)

    def _add_correlation_section(self, analysis_results: Dict[str, pd.DataFrame]):
        self._add_heading_cn("6  多指标相关性分析", level=1)
        self._add_paragraph_cn(
            f"本研究采用 Pearson 与 Spearman 两种方法计算代用指标之间的相关系数，"
            f"置信度水平 {int(self.config.get_correlation_settings().confidence_level * 100)}%。"
        )

        if self.settings.include_correlation_table and "pairwise_correlations" in analysis_results:
            pw = analysis_results["pairwise_correlations"].copy()
            if "p_value" in pw.columns:
                pw["显著性"] = pw["p_value"].apply(
                    lambda x: "***" if x < 0.001 else ("**" if x < 0.01 else ("*" if x < 0.05 else "n.s."))
                )
            self._add_dataframe_table(pw, title="表 4  代用指标两两相关性统计")

        self._add_heading_cn("6.1  相关性热力图", level=2)
        self._add_paragraph_cn("图 7～图 8 分别展示 Pearson 与 Spearman 相关系数矩阵热力图，色块颜色越深表示相关性越强。")
        for m in ["pearson", "spearman"]:
            key = f"heatmap_{m}"
            if key in analysis_results.get("correlation_matrix_" + m, {}):
                continue

    def _add_multi_core_comparison_section(
        self,
        comparison_results: Dict[str, pd.DataFrame],
        figure_paths: Dict[str, Dict[str, str]],
    ):
        self.doc.add_page_break()
        self._add_heading_cn("7  多钻孔地层对比分析", level=1)
        self._add_paragraph_cn(
            "本章对多个冰芯与石笋钻孔的代用指标数据进行跨区域对比分析，"
            "通过统一年代标尺下的差值计算与时空格局识别，揭示不同区域气候演化的同步性与差异性。"
        )

        self._add_heading_cn("7.1  跨钻孔 δ¹⁸O 时序对比", level=2)
        self._add_paragraph_cn(
            "图 9 展示了所有参与对比的钻孔在统一年代标尺下的 δ¹⁸O 时序曲线。"
            "各钻孔曲线沿同一时间轴排列，可直观比较不同区域气候事件的相位关系与幅度差异。"
        )

        if "multi_core_linked_view" in figure_paths:
            fig_paths = figure_paths["multi_core_linked_view"]
            if "png" in fig_paths:
                self._add_image_figure(fig_paths["png"], "图 9  多钻孔 δ¹⁸O 时序与差值联动视图")
            elif "html" in fig_paths:
                self._add_paragraph_cn(f"[交互式图表: {os.path.basename(fig_paths['html'])}]")

        if "epoch_wise_comparison" in comparison_results:
            self._add_heading_cn("7.2  各地层单位钻孔对比统计", level=2)
            self._add_paragraph_cn(
                "表 5 按地层单位统计了各钻孔 δ¹⁸O 的均值、标准差等特征，"
                "便于定量比较不同地质时期各区域的气候状态差异。"
            )
            epoch_df = comparison_results["epoch_wise_comparison"]
            self._add_dataframe_table(epoch_df, title="表 5  各地层单位-各钻孔 δ¹⁸O 统计对比")

        if "correlation_matrix_pearson" in comparison_results:
            self._add_heading_cn("7.3  跨钻孔相关性矩阵", level=2)
            self._add_paragraph_cn(
                "表 6 展示了各钻孔之间的 Pearson 相关系数矩阵，"
                "反映了不同区域气候信号在万年尺度上的协同变化程度。"
            )
            corr_df = comparison_results["correlation_matrix_pearson"].copy()
            self._add_dataframe_table(corr_df, title="表 6  跨钻孔 Pearson 相关系数矩阵")

        if "epoch_core_heatmap_mean" in figure_paths:
            self._add_heading_cn("7.4  地层-钻孔均值热力图", level=2)
            self._add_paragraph_cn(
                "图 10 以热力图形式直观展示了各地层单位中各钻孔 δ¹⁸O 均值的空间格局，"
                "颜色冷暖反映数值高低，便于识别区域分异规律。"
            )
            fig_paths = figure_paths.get("epoch_core_heatmap_mean", {})
            if "png" in fig_paths:
                self._add_image_figure(fig_paths["png"], "图 10  地层单位-钻孔均值热力图")
            elif "html" in fig_paths:
                self._add_paragraph_cn(f"[交互式图表: {os.path.basename(fig_paths['html'])}]")

        if "cross_core_diff_heatmap" in figure_paths:
            self._add_heading_cn("7.5  跨区域气候差异热力图", level=2)
            self._add_paragraph_cn(
                "图 11 为跨钻孔差值热力图，通过时间滑块可查看不同年代各钻孔之间的 δ¹⁸O 差异矩阵，"
                "揭示气候空间梯度随时间的演化特征。红色表示差值为正，蓝色表示差值为负。"
            )
            fig_paths = figure_paths.get("cross_core_diff_heatmap", {})
            if "html" in fig_paths:
                self._add_paragraph_cn(f"[交互式热力图: {os.path.basename(fig_paths['html'])}（可拖动年代滑块）]")
            if "png" in fig_paths:
                self._add_image_figure(fig_paths["png"], "图 11  跨钻孔差值热力图（最新年代快照）")

        self._add_paragraph_cn(
            "多钻孔对比分析结果表明，各区域气候演化在轨道尺度上具有显著的同步性，"
            "但在千年-百年尺度上存在明显的区域差异，反映了不同气候系统对外部强迫响应的敏感性差异。"
        )

    def _add_conclusion(self):
        self._add_heading_cn("8  结论与展望", level=1)
        self._add_paragraph_cn(
            "本研究通过冰芯与石笋多代用指标的联合分析，重建了研究区万年尺度的温度与降水演化序列，"
            "主要结论如下："
        )
        conclusions = [
            "（1）多源数据联合校正后的年代序列可有效识别末次冰盛期、新仙女木事件、全新世大暖期等关键气候事件。",
            "（2）冰芯 δ¹⁸O 反演温度与石笋 δ¹⁸O 反演降水在轨道-千年尺度上表现出显著的协变特征。",
            "（3）各地层单位代用指标分布存在显著差异，支持晚第四纪气候阶段性演化的经典认识。",
            "（4）建立的模块化数据处理框架支持新增地层数据源的增量重算，便于后续研究扩展。",
        ]
        for c in conclusions:
            self._add_paragraph_cn(c)
        self._add_paragraph_cn(
            "未来工作将进一步引入孢粉、湖相沉积等更多代用指标，并开展区域尺度的古气候数据合成与模型对比研究。"
        )

    def _add_appendix(self):
        self.doc.add_page_break()
        self._add_heading_cn("附录", level=1)
        self._add_heading_cn("附录 A  年代校正事件对照表", level=2)
        calib = self.config.get_age_calibration_table()
        if not calib.empty:
            self._add_dataframe_table(calib, title="表 A-1  主要气候事件年代校正对照表")
        self._add_heading_cn("附录 B  软件与版本", level=2)
        self._add_paragraph_cn("分析环境：Python 3.10+；主要库版本：Pandas 2.x、NumPy 1.x、Plotly 5.x、SciPy 1.x、python-docx 1.x")
